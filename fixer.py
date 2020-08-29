from PyQt5 import QtWidgets, QtGui, QtCore
import docx
import os
import sys

#doc = docx.Document(".//database.docx")
#new_doc = docx.Document()
#counter = 0
#for i in doc.paragraphs:
#    counter += 1
#    if i.text.strip().startswith("I") and "\n" in i.text:
#        new_i_list = i.text.strip().split("\n")
#        new_1 = new_i_list[0]
#        new_2 = new_i_list[1]
#        new_doc.add_paragraph(new_1)
#        new_doc.add_paragraph(new_2)
#    else:
#        new_doc.add_paragraph(i.text)

def goToWorkWindow(window):
    global wwindow 
    wwindow = WorkWindow(window.currentFile, window.keyss)
    window.hide()
    wwindow.show()

class WorkWindow(QtWidgets.QWidget):
    def __init__(self,currentFile,keyss, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.currentFile = currentFile
        self.keyss = keyss
        self.progressLabel = QtWidgets.QLabel("Ход выполнения: ")
        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.progressLabel)
        self.setLayout(self.vbox)

    def parse(self):
        print(1)


class FirstWindow(QtWidgets.QWidget):
    def __init__(self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.keyss = [] 

        #текстовое поле с путём
        self.label1 = QtWidgets.QLabel(str(os.getcwd()))

        #список с папками в директории скрипта
        self.listwid = QtWidgets.QListWidget()
        self.listdir = [i for i in os.listdir(".") if os.path.isdir("."+"//"+i)]#список директорий
        for i in self.listdir:
            if os.path.isdir("."+str("//")+i):
                self.listwid.addItem(i)

        self.choiceDirButton = QtWidgets.QPushButton('Выбрать')
        self.choiceDirButton.clicked.connect(self.choiceDir)

        self.subGrid = QtWidgets.QGridLayout()

        self.checkIBox  = QtWidgets.QCheckBox("Имеется ли 'I' метка?")
        self.subGrid.addWidget(self.checkIBox, 0, 0, 1, 2)

        self.startBoxLabel = QtWidgets.QLabel("Метка начала вопроса:")
        self.subGrid.addWidget(self.startBoxLabel, 1, 0)

        self.comboStartQuest = QtWidgets.QComboBox()
        listStartsPoints = ["S", "ЧИСЛО)", "Вопрос ЧИСЛО.", "ЧИСЛО."]
        self.comboStartQuest.addItems(listStartsPoints)
        self.subGrid.addWidget(self.comboStartQuest, 1, 1)

        self.questBoxLabel = QtWidgets.QLabel("Определение верности ответа:")
        self.subGrid.addWidget(self.questBoxLabel, 2, 0)

        self.comboRightQuest = QtWidgets.QComboBox()
        listRightPoints = ["(+/-):", "цвет текста - красный", "выделение, цвет - зеленый", "выделение, цвет - желтый", "нижнее подчеркивание", "курсив"]
        self.comboRightQuest.addItems(listRightPoints)
        self.subGrid.addWidget(self.comboRightQuest, 2, 1)


        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.label1)
        self.vbox.addWidget(self.listwid)
        self.vbox.addLayout(self.subGrid)
        self.vbox.addWidget(self.choiceDirButton)
        self.setLayout(self.vbox)

    def choiceDir(self):
        print('.'+"/"+self.listdir[self.listwid.currentRow()])
        self.currentFile = self.listdir[self.listwid.currentRow()]
        iPoint = self.checkIBox.checkState()
        startQuest = self.comboStartQuest.currentText()
        rightQuest = self.comboRightQuest.currentText()
        self.keyss = [iPoint, startQuest, rightQuest]
        print(self.keyss)

        goToWorkWindow(self)

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = FirstWindow()
    window.setWindowTitle("Котелина")
    window.resize(300, 200)
    desktop = QtWidgets.QApplication.desktop()
    x = (desktop.width() - window.width()) // 2
    y = (desktop.height() - window.height()) // 2
    window.move(x, y)
    window.show()
    sys.exit(app.exec_())


new_doc.save("fixed_database.docx")
