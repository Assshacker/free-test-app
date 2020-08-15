from PyQt5 import QtWidgets, QtCore, QtGui
import sys
import os
import docx
import time
import traceback
import random

#класс первоначального окна,в котором выбирается папка с тестом внутри,сами папки должны лежать
#в той же директории, что и скрипт
#hello git
#Q(queque)-выстроить правильный порядок вариантов ответа
#S(simple)-выбрать один правильный ответ
#C(conformity)-соотнести варианты ответов
#O(open)-открытый вопрос

#----------функции перехода между окнами-------------------------------------

def EndOfTesting(lisres):
    global reswin
    reswin = ResultWindow(lisres)
    window2.hide()
    reswin.show()

def FromResultToMain():
    reswin.hide()
    window.show()


def ModalWindowForward(window1):
    global window2
    window2 = SecondWindow(window1.currentFile)
    window2.getCountQuestion(int(window1.questCounter.text()))
    window2.resize(300, 200)
    window1.hide()
    window2.show()

def ModalWindowBackward():
    window2.hide()
    window.show()

#-----------------------------------------------------------------------

class TroubleWindow(QtWidgets.QWidget):
    def __init__(self,trLogList, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.trLogList = trLogList
        trLog = self.trLogList[0]
        problemText = self.trLogList[1]

        self.trLabelText = QtWidgets.QLabel(problemText)
        self.trLabel = QtWidgets.QLabel(trLog)

        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.trLabelText)
        self.vbox.addWidget(self.trLabel)

        self.setLayout(self.vbox)

    def checkAnswer(self):
        return True


class Database():
    def __init__(self, file_docx):
        self.file_docx = file_docx
        self.bankOfQuestions = []
        self.counter = 0
        self.doc = None
        self.status = None

    def load_file(self):
        try:
            self.doc = docx.Document("./"+self.file_docx+"/database.docx")
            print("База данных загружена успешно.")
            for i in self.doc.paragraphs:
                self.counter += 1
                if "I:" in i.text:
                    self.bankOfQuestions.append(self.counter)
            self.status = 0 
        except:
            print("Ошибка загрузки базы данных")
            print(traceback.format_exc())
            self.status = 1 

    def parser(self):
        try:
            currentTask = {}
            #qline = random.choice(self.bankOfQuestions)#строчка с "S:"
            qline = 1184
            startqline = qline
            print("случайная строчка: "+str(qline))
            currentTask["key"] = self.doc.paragraphs[qline].text[0]#присваиваем значение вначале строки: S, Q, O, и тд
            currentTask["question"] = self.doc.paragraphs[qline].text[2:]#призваиваем вопрос
            currentTask["answers"] = {}
            qline += 1
            if currentTask["key"] == "S" or currentTask["key"] == "Q":
                startchar = self.doc.paragraphs[qline].text.strip()[0]
                if startchar == "R" or startchar == "L":
                    currentTask["key"] = "C"
                    while (not ("I:" in self.doc.paragraphs[qline].text)) and (not (self.doc.paragraphs[qline].text.isspace())) and len(self.doc.paragraphs[qline].text) >= 4:
                        right, answ= self.doc.paragraphs[qline].text.split(":", maxsplit=1)
                        currentTask["answers"][right.strip()] = answ.strip()
                        qline += 1
                else: 
                    #while (not ("I:" in self.doc.paragraphs[qline].text)) and (not (self.doc.paragraphs[qline].text.isspace())) and len(self.doc.paragraphs[qline].text) >= 4:
                        #print(self.doc.paragraphs[qline].text)
                        #qline += 1
                    while (not ("I:" in self.doc.paragraphs[qline].text)) and (not (self.doc.paragraphs[qline].text.isspace())) and len(self.doc.paragraphs[qline].text) >= 4:
                        right, answ= self.doc.paragraphs[qline].text.split(":", maxsplit=1)
                        currentTask["answers"][answ.strip()] = right.strip()
                        qline += 1
                print(currentTask)
                return 0, currentTask
            elif currentTask["key"] == "C":
                print(1)
            elif currentTask["key"] == "O":
                print(2)
        except:
            trLog = traceback.format_exc()
            textProblem = " "
            qline = startqline
            while (not ("I:" in self.doc.paragraphs[qline].text)) and (not (self.doc.paragraphs[qline].text.isspace())) and len(self.doc.paragraphs[qline].text) >= 4:
                textProblem = textProblem + self.doc.paragraphs[qline].text + "\n"
                qline += 1
            return 1, [trLog, textProblem]


class ResultWindow(QtWidgets.QWidget):
    def __init__(self,result, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.result = result
        self.labelResult = QtWidgets.QLabel("   Вы решили {}  вопросов, из которых:\n   {} - правильных ({} %)\n   {} - неправильных ({} %) ".format(self.result[0], self.result[1], (self.result[1]/self.result[0])*100, self.result[2], (self.result[2]/self.result[0])*100))
        self.cancelButton = QtWidgets.QPushButton("В меню")
        self.cancelButton.clicked.connect(self.exit)
        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.labelResult)
        self.vbox.addWidget(self.cancelButton)
        self.setLayout(self.vbox)

    def exit(self):
        FromResultToMain()

class SingleQuestionWindow(QtWidgets.QWidget):
    def __init__(self,currentTask, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.currentTask = currentTask
        self.questLabel = QtWidgets.QLabel(self.currentTask["question"])
        
        self.groupBox = QtWidgets.QGroupBox("")
        self.vBoxGroupBox = QtWidgets.QVBoxLayout()
        lencq = len(self.currentTask["answers"])
        self.listBut = []
        if lencq == 1:
            self.but1 = QtWidgets.QRadioButton("1")
            self.listBut.append(self.but1)
        elif lencq == 2:
            self.but1 = QtWidgets.QRadioButton("1")
            self.but2 = QtWidgets.QRadioButton("2")
            self.listBut.append(self.but2)
            self.listBut.append(self.but1)
        elif lencq == 3:
            self.but1 = QtWidgets.QRadioButton("1")
            self.but2 = QtWidgets.QRadioButton("2")
            self.but3 = QtWidgets.QRadioButton("3")
            self.listBut.append(self.but1)
            self.listBut.append(self.but2)
            self.listBut.append(self.but3)
        elif lencq == 4:
            self.but1 = QtWidgets.QRadioButton("1")
            self.but2 = QtWidgets.QRadioButton("2")
            self.but3 = QtWidgets.QRadioButton("3")
            self.but4 = QtWidgets.QRadioButton("4")
            self.listBut.append(self.but1)
            self.listBut.append(self.but2)
            self.listBut.append(self.but3)
            self.listBut.append(self.but4)
        elif lencq  == 5:
            self.but1 = QtWidgets.QRadioButton("1")
            self.but2 = QtWidgets.QRadioButton("2")
            self.but3 = QtWidgets.QRadioButton("3")
            self.but4 = QtWidgets.QRadioButton("4")
            self.but5 = QtWidgets.QRadioButton("5")
            self.listBut.append(self.but1)
            self.listBut.append(self.but2)
            self.listBut.append(self.but3)
            self.listBut.append(self.but4)
            self.listBut.append(self.but5)
        print(self.listBut)

        self.answersFromKeys = list(self.currentTask["answers"].keys()) 

        for i in range(len(self.listBut)):
            self.listBut[i].setText(self.answersFromKeys[i])

        random.shuffle(self.listBut)

        for i in self.listBut:
            self.vBoxGroupBox.addWidget(i) 
        self.groupBox.setLayout(self.vBoxGroupBox)

        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.questLabel)
        self.vbox.addWidget(self.groupBox)

        self.setLayout(self.vbox)

    def checkAnswer(self):
        for i in self.listBut:
            if i.isChecked() and self.currentTask["answers"][i.text()] == "+":
                return True
            elif i.isChecked() and self.currentTask["answers"][i.text()] == "-":
                return False

class QuequeQuestionWindow(QtWidgets.QWidget):
    def __init__(self,currentTask, parent = None):
        QtWidgets.QWidget.__init__(self, parent)
        self.currentTask = currentTask
        self.questLabel = QtWidgets.QLabel(self.currentTask["question"]+"\n(первый ответ - вверху, последний - внизу)")
        self.qComboList = []
        lencq = len(self.currentTask["answers"])  
        if lencq == 2:
            self.qcombo1 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo1)

            self.qcombo2 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo2)

        elif lencq == 3:
            self.qcombo1 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo1)

            self.qcombo2 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo2)
            
            self.qcombo3 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo3)

        elif lencq == 4:
            self.qcombo1 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo1)

            self.qcombo2 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo2)
            
            self.qcombo3 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo3)

            self.qcombo4 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo4)

        elif lencq == 5:
            self.qcombo1 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo1)

            self.qcombo2 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo2)
            
            self.qcombo3 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo3)

            self.qcombo4 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo4)
            
            self.qcombo5 = QtWidgets.QComboBox()
            self.qComboList.append(self.qcombo5)

        self.answersFromKeys = list(self.currentTask["answers"].keys())
        random.shuffle(self.answersFromKeys)
        self.answersFromKeys.insert(0, "Выберите ответ")
        
        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.questLabel)

        for i in range(len(self.qComboList)):
            self.qComboList[i].addItems(self.answersFromKeys)
            self.vbox.addWidget(self.qComboList[i])

        self.setLayout(self.vbox)
    def checkAnswer(self):
        for i in self.qComboList:
            if (self.qComboList.index(i) + 1) == int(self.currentTask["answers"][i.currentText()]) :
                continue
            else:
                return False
        return True

class ConformityQuestionWindow(QtWidgets.QWidget):
    def __init__(self, currentTask,  parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.currentTask = currentTask

        Llist = []
        Rlist =[]
        LlistEqual = {}
        RlistEqual = {}
        maxList = None
        self.staticAnswers = []
        self.dinamicAnswers = []
        self.resultSlov = {}
        minKey = None

        self.questLabel = QtWidgets.QLabel(self.currentTask["question"])
        self.mainGrid = QtWidgets.QGridLayout()
        self.mainGrid.addWidget(self.questLabel, 0, 0, 1, 2)

        for i in list(self.currentTask["answers"].items()):
            if i[0][0] == "L":
                Llist.append(i)
            else:
                Rlist.append(i)

        for i in Llist:
            for a in Llist:
                if i == a:
                    continue
                else:
                    if i[1] == a[1]:
                        LlistEqual[i[0]] = a[0]

        for i in Rlist:
            for a in Rlist:
                if i == a:
                    continue
                else:
                    if i[1] == a[1]:
                        RlistEqual[i[0]] = a[0]

        if len(LlistEqual) == 0:
            minKey = "L"
            for i in range(len(Llist)):
                self.resultSlov[Llist[i][1]] = Rlist[i][1]
        elif len(RlistEqual) == 0:
            minKey = "R"
            for i in range(len(Rlist)):
                self.resultSlov[Rlist[i][1]] = Llist[i][1]

        raz = len(Llist) - len(Rlist)
        if raz < 0:
            for i in range(1, (abs(raz)+1)):
                self.resultSlov[Rlist[-i][1]] = "w"+str(i)
                
        #print(self.resultSlov)

        Llist.clear() #теперь список для значений для статичных текстовых полей
        Rlist.clear()#а это значения для виджета выборки

        for i in list(self.resultSlov.values()):
            if not i.startswith("w"):
                Llist.append(i)
        
        for i in list(self.resultSlov.keys()):
            Rlist.append(i)

        random.shuffle(Llist)
        random.shuffle(Rlist)
        Rlist.insert(0, "Выберите ответ")

        cs = len(Llist)

        if cs == 1:
            self.answLabel1 = QtWidgets.QLabel("1")
            self.staticAnswers.append(self.answLabel1)
        elif cs == 2:
            self.answLabel1 = QtWidgets.QLabel("1")
            self.staticAnswers.append(self.answLabel1)
            self.answLabel2 = QtWidgets.QLabel("2")
            self.staticAnswers.append(self.answLabel2)
        elif cs == 3:
            self.answLabel1 = QtWidgets.QLabel("1")
            self.staticAnswers.append(self.answLabel1)
            self.answLabel2 = QtWidgets.QLabel("2")
            self.staticAnswers.append(self.answLabel2)
            self.answLabel3 = QtWidgets.QLabel("3")
            self.staticAnswers.append(self.answLabel3)
        elif cs == 4:
            self.answLabel1 = QtWidgets.QLabel("1")
            self.staticAnswers.append(self.answLabel1)
            self.answLabel2 = QtWidgets.QLabel("2")
            self.staticAnswers.append(self.answLabel2)
            self.answLabel3 = QtWidgets.QLabel("3")
            self.staticAnswers.append(self.answLabel3)
            self.answLabel4 = QtWidgets.QLabel("4")
            self.staticAnswers.append(self.answLabel4)
        elif cs == 5:
            self.answLabel1 = QtWidgets.QLabel("1")
            self.staticAnswers.append(self.answLabel1)
            self.answLabel2 = QtWidgets.QLabel("2")
            self.staticAnswers.append(self.answLabel2)
            self.answLabel3 = QtWidgets.QLabel("3")
            self.staticAnswers.append(self.answLabel3)
            self.answLabel4 = QtWidgets.QLabel("4")
            self.staticAnswers.append(self.answLabel4)
            self.answLabel5 = QtWidgets.QLabel("5")
            self.staticAnswers.append(self.answLabel5)
        elif cs == 6:
            self.answLabel1 = QtWidgets.QLabel("1")
            self.staticAnswers.append(self.answLabel1)
            self.answLabel2 = QtWidgets.QLabel("2")
            self.staticAnswers.append(self.answLabel2)
            self.answLabel3 = QtWidgets.QLabel("3")
            self.staticAnswers.append(self.answLabel3)
            self.answLabel4 = QtWidgets.QLabel("4")
            self.staticAnswers.append(self.answLabel4)
            self.answLabel5 = QtWidgets.QLabel("5")
            self.staticAnswers.append(self.answLabel5)
            self.answLabel6 = QtWidgets.QLabel("6")
            self.staticAnswers.append(self.answLabel6)

        for i in range(len(self.staticAnswers)):
            self.staticAnswers[i].setText(Llist[i])

        if cs == 1:
            self.answBox1 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox1)
        elif cs == 2:
            self.answBox1 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox1)
            self.answBox2 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox2)
        elif cs == 3:
            self.answBox1 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox1)
            self.answBox2 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox2)
            self.answBox3 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox3)
        elif cs == 4:
            self.answBox1 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox1)
            self.answBox2 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox2)
            self.answBox3 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox3)
            self.answBox4 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox4)
        elif cs == 5:
            self.answBox1 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox1)
            self.answBox2 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox2)
            self.answBox3 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox3)
            self.answBox4 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox4)
            self.answBox5 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox5)
        elif cs == 6:
            self.answBox1 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox1)
            self.answBox2 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox2)
            self.answBox3 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox3)
            self.answBox4 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox4)
            self.answBox5 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox5)
            self.answBox6 = QtWidgets.QComboBox()
            self.dinamicAnswers.append(self.answBox6)

        for i in self.dinamicAnswers:
            i.addItems(Rlist)

        for i in range(len(self.staticAnswers)):
            self.mainGrid.addWidget(self.staticAnswers[i],i+1,0)
            self.mainGrid.addWidget(self.dinamicAnswers[i],i+1,1)

        self.setLayout(self.mainGrid)

    def checkAnswer(self):
        for i in range(len(self.staticAnswers)):
            if self.resultSlov[self.dinamicAnswers[i].currentText()] == self.staticAnswers[i].text():
                continue
            else:
                return False
        return True


class FirstWindow(QtWidgets.QWidget):
    def __init__(self, parent=None):
        QtWidgets.QWidget.__init__(self, parent)

        #текстовое поле с путём
        self.label1 = QtWidgets.QLabel(str(os.getcwd()))
        
        #список с папками в директории скрипта
        self.listwid = QtWidgets.QListWidget()
        self.listdir = [i for i in os.listdir(".") if os.path.isdir("."+"//"+i)]#список директорий
        for i in self.listdir:
            if os.path.isdir("."+str("//")+i):
                self.listwid.addItem(i)

        self.choiceDirButton = QtWidgets.QPushButton('Выбрать')
        self.choiceDirButton.clicked.connect(self.choiceDir)

        #область выбора кол-ва вопросов: слева надпись справа счетчик
        self.questCounter = QtWidgets.QSpinBox()
        self.questCounter.setMinimum(1)
        self.questCounter.setValue(40)
        self.labelQuestCounter = QtWidgets.QLabel("Кол-во вопросов:")
        self.subhbox = QtWidgets.QHBoxLayout()
        self.subhbox.addWidget(self.labelQuestCounter)
        self.subhbox.addWidget(self.questCounter)

        self.vbox = QtWidgets.QVBoxLayout()
        self.vbox.addWidget(self.label1)
        self.vbox.addWidget(self.listwid)
        self.vbox.addLayout(self.subhbox)
        self.vbox.addWidget(self.choiceDirButton)
        self.setLayout(self.vbox)

    def choiceDir(self):
        print('.'+"/"+self.listdir[self.listwid.currentRow()])
        self.currentFile = self.listdir[self.listwid.currentRow()]
        ModalWindowForward(self)


class SecondWindow(QtWidgets.QWidget):
    def __init__(self, currentFile, parent=None):
        QtWidgets.QWidget.__init__(self, parent)
        self.currentFile = currentFile 
        self.questCounter = 0
        self.rightAnswers = 0
        self.wrongAnswers = 0
        self.currentQuestion = 1
        self.isAnswered = False

        self.vbox = QtWidgets.QVBoxLayout()
        self.database = Database(currentFile)
        self.database.load_file()
        
        
        if self.database.status == 1:
            self.label1 = QtWidgets.QLabel("ошибка базы данных")
            self.vbox.addWidget(self.label1)

            self.cancelButton = QtWidgets.QPushButton("Cancel")
            self.cancelButton.clicked.connect(self.cancel)
            self.vbox.addWidget(self.cancelButton)

        elif self.database.status == 0:
            trStatus, self.task = self.database.parser()
            print(trStatus)
            if trStatus == 1:
                self.taskWindow = TroubleWindow(self.task)
            else:
                if self.task["key"] == "S":
                    self.taskWindow = SingleQuestionWindow(self.task)
                elif self.task["key"] == "Q":
                    self.taskWindow = QuequeQuestionWindow(self.task)
                elif self.task["key"] == "C":
                    self.taskWindow = ConformityQuestionWindow(self.task)

            self.cancelButton = QtWidgets.QPushButton("Выход")
            self.cancelButton.clicked.connect(self.cancel)

            self.statusCountLabel = QtWidgets.QLabel(" ")

            self.checkButton = QtWidgets.QPushButton("Проверить")
            self.checkButton.clicked.connect(self.check)

            self.nextQuestionButton = QtWidgets.QPushButton("Далее")
            self.nextQuestionButton.clicked.connect(self.nextQuestion)

            self.subhbox = QtWidgets.QHBoxLayout()
            self.subhbox.addWidget(self.cancelButton)
            self.subhbox.addWidget(self.nextQuestionButton)


            self.vbox.addWidget(self.statusCountLabel)
            self.vbox.addWidget(self.taskWindow)
            self.vbox.addWidget(self.checkButton)
            self.vbox.addLayout(self.subhbox)

        self.setLayout(self.vbox)

    def cancel(self):
        ModalWindowBackward()

    def updateText(self):
        self.statusCountLabel.setText("%d из %d, %d верных, %d неверных" % (self.currentQuestion, self.questCounter, self.rightAnswers, self.wrongAnswers))

    def getCountQuestion(self, countQuestion):
        if self.database.status == 0:
            self.questCounter = countQuestion
            self.updateText()

    def check(self):
        answ = self.taskWindow.checkAnswer()
        if answ == True:
            self.checkButton.setText("Правильно")
            if self.isAnswered == False:
                self.rightAnswers += 1
                self.isAnswered = True
        elif answ == False:
            self.checkButton.setText("Неверно")
            if self.isAnswered == False:
                self.wrongAnswers += 1
                self.isAnswered = True
        self.updateText()

    def nextQuestion(self):
        if self.currentQuestion < self.questCounter:
            self.isAnswered = False
            self.checkButton.setText("Проверить")
            self.currentQuestion += 1
            self.task = self.database.parser()
            self.vbox.removeWidget(self.taskWindow)
            self.taskWindow.setParent(None)
            self.taskWindow = None
            self.taskWindow = SingleQuestionWindow(self.task)
            self.vbox.insertWidget(1, self.taskWindow)
            self.updateText()
        else:
            self.endOfTest()
    
    def endOfTest(self):
        listResult = [self.questCounter, self.rightAnswers, self.wrongAnswers]
        print(listResult)
        EndOfTesting(listResult)



if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = FirstWindow()
    window.setWindowTitle("testshmest")
    window.resize(300, 200)
    window.show()
    sys.exit(app.exec_())
       
#startchar = self.doc.paragraphs[qline].text.strip()[0]
#            
#                while (not ("I:" in self.doc.paragraphs[qline].text)) and (not (self.doc.paragraphs[qline].text.isspace())) and len(self.doc.paragraphs[qline].text) >= 4:
#            else:
